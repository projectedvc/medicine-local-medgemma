import html
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.core.config import settings
from app.models.ai import AIAnalysis
from app.models.enums import FindingClass
from app.models.report import Report
from app.models.study import Study
from app.models.user import User
from app.services.report_localization import build_localized_ai_draft, finding_label, normalize_report_lang
from app.services.translation import translate_text


DISCLAIMER = (
    "AI нәтижесі алдын ала клиникалық көмек ретінде ғана қолданылады. "
    "Қорытынды шешімді дәрігер қабылдайды. / "
    "Результат AI является предварительной подсказкой. "
    "Окончательное решение принимает врач."
)

REPORT_TITLE = "Медициналық қорытынды / Медицинское заключение"
REPORT_SUBTITLE = "MedAI Radiology Assistant"

FINDING_LABELS: dict[FindingClass, str] = {
    FindingClass.normal: "Норма / Normal",
    FindingClass.pneumonia: "Пневмония / Pneumonia",
    FindingClass.other_abnormal: "Басқа патология / Другая патология",
    FindingClass.pleural_effusion: "Плевралық сұйықтық / Плевральный выпот",
    FindingClass.pneumothorax: "Пневмоторакс / Pneumothorax",
    FindingClass.atelectasis: "Ателектаз / Atelectasis",
}

TEMPLATES: dict[FindingClass, str] = {
    FindingClass.normal: (
        "Айқын жедел кардиопульмоналдық өзгеріс анықталмады. Өкпе алаңдары айқын, "
        "плевралық сұйықтық немесе пневмоторакс белгілері көрінбейді."
    ),
    FindingClass.pneumonia: (
        "Өкпе тінінде инфильтративті өзгерістерге сәйкес болуы мүмкін белгілер байқалады. "
        "Клиникалық деректермен және дәрігердің тікелей бағалауымен салыстыру қажет."
    ),
    FindingClass.other_abnormal: (
        "Кеуде рентгенограммасында пневмония ретінде сенімді жіктелмейтін патологиялық өзгерістер бар. "
        "Суретті рентгенолог тікелей сипаттауы қажет."
    ),
    FindingClass.pleural_effusion: (
        "Плевра қуысында сұйықтық болуы мүмкін белгілер байқалады. Жағы, көлемі және "
        "клиникалық маңызы дәрігермен нақтылануы тиіс."
    ),
    FindingClass.pneumothorax: (
        "Пневмотораксқа күмәнді белгілер байқалады. Науқастың жағдайын және суретті "
        "дәрігер шұғыл бағалауы қажет."
    ),
    FindingClass.atelectasis: (
        "Өкпе көлемінің жергілікті төмендеуі/ателектаз болуы мүмкін белгілер байқалады. "
        "Қорытынды дәрігерлік верификацияны талап етеді."
    ),
}

SECTION_ALIASES = {
    "findings": "Көрініс / Описание",
    "impression": "Қорытынды / Заключение",
    "recommendations": "Ұсыныс / Рекомендация",
}


def _project_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _stamp_path() -> Path | None:
    path = _project_root() / "pec.png"
    return path if path.exists() else None


def _latest_image_path(study: Study) -> Path | None:
    if not study.images:
        return None
    image = study.images[-1]
    path = Path(image.preview_path or image.storage_path)
    return path if path.exists() else None


def _clean_line(line: str) -> str:
    line = html.unescape(line).strip()
    line = re.sub(r"<[^>]+>", " ", line)
    line = re.sub(r"^[\-\*\d\.\)\s]+", "", line)
    line = re.sub(r"\s+", " ", line)
    return line.strip()


def _strip_reasoning(text: str) -> str:
    text = html.unescape(text or "")
    text = re.sub(r"(?i)\bmedgemma\b", "AI", text)
    text = re.sub(r"<unused\d+>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</?s>|<bos>|<eos>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"\r\n?", "\n", text)

    lines: list[str] = []
    skip_block = False
    skip_markers = (
        "thought",
        "we need answer",
        "the user wants",
        "the prompt asks",
        "i should",
        "we should",
        "analyze this image",
        "analyze the image",
        "consider the clinical context",
        "provide a json",
        "json output",
        "json",
        "schema",
        "the task",
        "system instruction",
        "assistant answer",
        "assistant:",
        "user:",
        "орк",
        "ogk",
        "this is a type for",
    )
    stop_skip_markers = (
        "findings",
        "impression",
        "recommendation",
        "recommendations",
        "conclusion",
        "diagnosis",
        "response",
        "prediction",
    )
    for raw_line in text.splitlines():
        line = _clean_line(raw_line)
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            continue

        lower = line.casefold()
        if any(lower.startswith(marker) for marker in stop_skip_markers):
            skip_block = False

        if skip_block:
            if any(marker in lower for marker in stop_skip_markers):
                skip_block = False
            else:
                continue

        if any(marker in lower for marker in skip_markers):
            skip_block = True
            continue

        if re.fullmatch(r"[{}\[\],:\"]+", line):
            continue
        lines.append(line)

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def _section_from_text(text: str, heading: str, next_headings: list[str]) -> str | None:
    pattern = rf"{re.escape(heading)}\s*:?\s*(.*?)"
    if next_headings:
        pattern += rf"(?=\n\s*(?:{'|'.join(map(re.escape, next_headings))})\s*:|\Z)"
    else:
        pattern += r"\Z"
    match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        return None
    value = _strip_reasoning(match.group(1))
    return value or None


def _parse_payload(raw_response_json: str | None) -> dict[str, Any]:
    if not raw_response_json:
        return {}
    try:
        payload = json.loads(raw_response_json)
    except ValueError:
        return {}
    return payload if isinstance(payload, dict) else {}


def _payload_section(payload: dict[str, Any], key: str) -> str | None:
    value = payload.get(key)
    if isinstance(value, str):
        value = _strip_reasoning(value)
        return value or None
    if isinstance(value, list):
        lines = [_clean_line(str(item)) for item in value if _clean_line(str(item))]
        return "\n".join(lines) if lines else None
    return None


def _generated_ai_sections(analysis: AIAnalysis | None) -> list[tuple[str, str]]:
    if not analysis:
        return []
    payload = _parse_payload(analysis.raw_response_json)
    if not payload:
        return []

    sections: list[tuple[str, str]] = []
    for key in ("findings", "impression", "recommendations"):
        value = _payload_section(payload, key)
        if value:
            sections.append((SECTION_ALIASES[key], value))

    if sections:
        return sections

    response = payload.get("response")
    if not isinstance(response, str) or not response.strip():
        return []

    response = _strip_reasoning(response)
    findings = _section_from_text(response, "Findings", ["Impression", "Recommendations", "Conclusion"])
    impression = _section_from_text(response, "Impression", ["Recommendations", "Conclusion"])
    recommendations = _section_from_text(response, "Recommendations", ["Conclusion"])

    if findings:
        sections.append((SECTION_ALIASES["findings"], findings))
    if impression:
        sections.append((SECTION_ALIASES["impression"], impression))
    if recommendations:
        sections.append((SECTION_ALIASES["recommendations"], recommendations))
    if not sections and response:
        sections.append(("MedAI", response))
    return sections


def _generated_ai_text(analysis: AIAnalysis) -> str | None:
    sections = _generated_ai_sections(analysis)
    if not sections:
        return None
    return "\n\n".join(f"{title}:\n{text}" for title, text in sections)


def _analysis_summary(analysis: AIAnalysis | None) -> list[str]:
    if not analysis or analysis.status.value != "completed":
        return ["AI талдау аяқталмаған немесе қолжетімсіз."]

    lines: list[str] = []
    generated = _generated_ai_text(analysis)
    if generated:
        lines.extend(["MedAI жауабы:", generated, ""])

    if analysis.hidden_due_low_confidence or not analysis.predicted_class:
        lines.extend(
            [
                "AI сенімділігі белгіленген шектен төмен.",
                "Диагностикалық класс жасырылды, дәрігердің қолмен бағалауы қажет.",
            ]
        )
        return lines

    label = FINDING_LABELS.get(analysis.predicted_class, analysis.predicted_class.value)
    confidence = f"{(analysis.confidence or 0) * 100:.1f}%"
    lines.extend(
        [
            f"AI алдын ала класы: {label}.",
            f"AI сенімділігі: {confidence}.",
            TEMPLATES[analysis.predicted_class],
        ]
    )
    return lines


def build_ai_draft(study: Study, analysis: AIAnalysis | None) -> str:
    lines = [
        "AI қорытынды жобасы / AI-черновик заключения",
        f"Зерттеу / Исследование: {study.accession_number}",
        f"Пациент коды / Код пациента: {study.patient_code}",
        f"Зерттеу түрі / Тип исследования: {study.study_type}",
        "",
    ]
    if study.clinical_note:
        lines.extend(["Клиникалық жазба / Клиническая заметка:", study.clinical_note, ""])
    lines.extend(_analysis_summary(analysis))
    lines.extend(
        [
            "",
            "Дәрігер финал мәтінді тексеріп, түзетіп және қолмен растауы керек.",
            "Врач должен проверить, отредактировать и вручную подтвердить финальное заключение.",
        ]
    )
    return "\n".join(lines)


def build_ai_draft(study: Study, analysis: AIAnalysis | None, lang: str | None = "ru") -> str:
    return build_localized_ai_draft(study, analysis, lang)


def ensure_export_dir() -> Path:
    settings.export_dir.mkdir(parents=True, exist_ok=True)
    return settings.export_dir


def _font_candidates() -> list[str]:
    candidates = []
    if settings.pdf_font_path:
        candidates.append(settings.pdf_font_path)
    candidates.extend(
        [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibri.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        ]
    )
    return candidates


def _initials(name: str) -> str:
    parts = [part for part in name.replace(".", " ").split() if part]
    if not parts:
        return "DR"
    return "".join(part[0].upper() for part in parts[:2])


def _doctor_photo(doctor: User, export_dir: Path) -> Path:
    from PIL import Image, ImageDraw, ImageFont

    path = export_dir / f"doctor-{doctor.id}-photo.png"
    if path.exists():
        return path

    image = Image.new("RGB", (420, 520), "#f5f8fa")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((18, 18, 402, 502), radius=28, outline="#c8d7df", width=4)
    draw.ellipse((126, 72, 294, 240), fill="#dbeaf0", outline="#8fb2c2", width=3)
    draw.rounded_rectangle((92, 270, 328, 430), radius=48, fill="#cfe0e8")

    font_path = next((candidate for candidate in _font_candidates() if Path(candidate).exists()), None)
    font_big = ImageFont.truetype(font_path, 76) if font_path else ImageFont.load_default()
    font_small = ImageFont.truetype(font_path, 22) if font_path else ImageFont.load_default()
    initials = _initials(doctor.full_name)
    bbox = draw.textbbox((0, 0), initials, font=font_big)
    draw.text(((420 - (bbox[2] - bbox[0])) / 2, 112), initials, fill="#27495a", font=font_big)
    draw.text((128, 454), "Doctor photo", fill="#5c7582", font=font_small)
    image.save(path)
    return path


def _format_dt(value: datetime | None = None) -> str:
    moment = value or datetime.now(timezone.utc)
    return moment.astimezone().strftime("%d.%m.%Y %H:%M")


def _analysis_for_study(study: Study) -> AIAnalysis | None:
    analyses = [analysis for analysis in study.analyses if analysis.completed_at or analysis.created_at]
    if not analyses:
        return None
    return sorted(analyses, key=lambda item: item.completed_at or item.created_at, reverse=True)[0]


def _status_chip(analysis: AIAnalysis | None) -> tuple[str, str]:
    if not analysis or analysis.status.value != "completed":
        return "AI талдау жоқ / AI анализ отсутствует", "#64748b"
    if analysis.hidden_due_low_confidence or not analysis.predicted_class:
        return "Төмен сенімділік / Низкая уверенность", "#b45309"
    label = FINDING_LABELS.get(analysis.predicted_class, analysis.predicted_class.value)
    confidence = f"{(analysis.confidence or 0) * 100:.1f}%"
    return f"{label} · {confidence}", "#047857"


def _status_chip(analysis: AIAnalysis | None, lang: str | None = "ru") -> tuple[str, str]:
    report_lang = normalize_report_lang(lang)
    if not analysis or analysis.status.value != "completed":
        fallback = {
            "kk": "AI талдау жоқ",
            "ru": "AI-анализ отсутствует",
            "en": "AI analysis unavailable",
        }
        return fallback[report_lang], "#64748b"
    if analysis.hidden_due_low_confidence or not analysis.predicted_class:
        fallback = {
            "kk": "Төмен сенімділік",
            "ru": "Низкая уверенность",
            "en": "Low confidence",
        }
        return fallback[report_lang], "#b45309"
    confidence = f"{(analysis.confidence or 0) * 100:.1f}%"
    return f"{finding_label(analysis.predicted_class, report_lang)} · {confidence}", "#047857"


def _clean_final_text(text: str | None, study: Study) -> list[tuple[str, str]]:
    cleaned = _strip_reasoning(text or "")
    if not cleaned:
        return [("Қорытынды / Заключение", "Қорытынды мәтіні енгізілмеген.")]

    remove_patterns = [
        r"^AI қорытынды жобасы\s*/\s*AI-черновик заключения\s*$",
        r"^Зерттеу\s*/\s*Исследование\s*:.*$",
        r"^Пациент коды\s*/\s*Код пациента\s*:.*$",
        r"^Зерттеу түрі\s*/\s*Тип исследования\s*:.*$",
        r"^Клиникалық жазба\s*/\s*Клиническая заметка\s*:?\s*$",
        r"^MedAI жауабы\s*:?\s*$",
        r"^MedAI local response\s*:?\s*$",
        r"^MedGemMA local response\s*:?\s*$",
        r"^AI сенімділігі белгіленген.*$",
        r"^Диагностикалық класс жасырылды.*$",
        r"^Дәрігер финал мәтінді.*$",
        r"^Врач должен проверить.*$",
        r"^Ескерту\s*/\s*Важно\s*:.*$",
        r"^ВАЖНО\..*$",
    ]
    filtered: list[str] = []
    note_seen = False
    for raw_line in cleaned.splitlines():
        line = _clean_line(raw_line)
        if not line:
            if filtered and filtered[-1] != "":
                filtered.append("")
            continue
        if any(re.match(pattern, line, flags=re.IGNORECASE) for pattern in remove_patterns):
            continue
        if study.clinical_note and line == _clean_line(study.clinical_note):
            if note_seen:
                continue
            note_seen = True
            continue
        filtered.append(line)

    cleaned = "\n".join(filtered)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        cleaned = "Қорытынды мәтіні дәрігермен толтырылады."

    sections: list[tuple[str, str]] = []
    known = [
        ("Сипаттама", "Сипаттама"),
        ("Қорытынды", "Қорытынды"),
        ("Ұсыныстар", "Ұсыныстар"),
        ("Описание", "Описание"),
        ("Заключение", "Заключение"),
        ("Рекомендации", "Рекомендации"),
        ("Көрініс / Описание", "Көрініс / Описание"),
        ("Описание", "Көрініс / Описание"),
        ("Findings", "Көрініс / Описание"),
        ("Қорытынды / Заключение", "Қорытынды / Заключение"),
        ("Заключение", "Қорытынды / Заключение"),
        ("Impression", "Қорытынды / Заключение"),
        ("Ұсыныс / Рекомендация", "Ұсыныс / Рекомендация"),
        ("Рекомендации", "Ұсыныс / Рекомендация"),
        ("Recommendations", "Ұсыныс / Рекомендация"),
    ]
    heading_pattern = "|".join(re.escape(item[0]) for item in known)
    parts = re.split(rf"(?im)^\s*({heading_pattern})\s*:?\s*$", cleaned)
    if len(parts) > 1:
        prefix = parts[0].strip()
        if prefix:
            sections.append(("Қорытынды / Заключение", prefix))
        title_map = {source.casefold(): target for source, target in known}
        for index in range(1, len(parts), 2):
            source_title = parts[index].strip().casefold()
            body = parts[index + 1].strip() if index + 1 < len(parts) else ""
            if body:
                sections.append((title_map.get(source_title, parts[index].strip()), body))
    else:
        sections.append(("Қорытынды / Заключение", cleaned))

    return sections


def _plain_paragraphs(text: str) -> list[str]:
    paragraphs: list[str] = []
    current: list[str] = []
    for line in text.splitlines():
        line = _clean_line(line)
        if not line:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue
        current.append(line)
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs or [""]


def _report_text_for_lang(study: Study, report: Report, analysis: AIAnalysis | None, lang: str) -> str:
    if analysis and analysis.status.value == "completed":
        return build_localized_ai_draft(study, analysis, lang)
    final_text = (report.final_text or "").strip()
    ai_draft = (report.ai_draft_text or "").strip()
    if analysis and (not final_text or final_text == ai_draft):
        return build_localized_ai_draft(study, analysis, lang)
    if not final_text:
        return build_localized_ai_draft(study, analysis, lang)
    translated = translate_text(final_text, lang) if lang in {"kk", "ru"} else None
    return translated or final_text


def _pdf_image(path: Path, max_width: float, max_height: float):
    from reportlab.platypus import Image as PdfImage

    image = PdfImage(str(path))
    ratio = min(max_width / image.imageWidth, max_height / image.imageHeight)
    image.drawWidth = image.imageWidth * ratio
    image.drawHeight = image.imageHeight * ratio
    return image


def _register_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for candidate in _font_candidates():
        if Path(candidate).exists():
            pdfmetrics.registerFont(TTFont("ClinicalSans", candidate))
            return "ClinicalSans"
    return "Helvetica"


def _pdf_para(text: str, style):
    from reportlab.platypus import Paragraph

    return Paragraph(html.escape(text).replace("\n", "<br/>"), style)


def export_report_pdf(study: Study, report: Report, doctor: User, lang: str | None = "ru") -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    export_dir = ensure_export_dir()
    path = export_dir / f"study-{study.id}-report.pdf"
    font_name = _register_pdf_font()
    lang = "kk"
    analysis = _analysis_for_study(study)
    chip_text, chip_color = _status_chip(analysis, lang)

    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "ClinicalNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#17232b"),
    )
    muted = ParagraphStyle(
        "ClinicalMuted",
        parent=normal,
        fontSize=7.5,
        leading=9.5,
        textColor=colors.HexColor("#667985"),
    )
    title = ParagraphStyle(
        "ClinicalTitle",
        parent=normal,
        fontName=font_name,
        fontSize=15,
        leading=18,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0b1f2a"),
        spaceAfter=2,
    )
    subtitle = ParagraphStyle(
        "ClinicalSubtitle",
        parent=muted,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    section = ParagraphStyle(
        "ClinicalSection",
        parent=normal,
        fontName=font_name,
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#0b6476"),
        spaceBefore=8,
        spaceAfter=5,
    )
    small_caps = ParagraphStyle(
        "ClinicalSmallCaps",
        parent=muted,
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor("#64748b"),
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=f"Clinical report {study.accession_number}",
    )

    story = [
        Paragraph("Медициналық қорытынды", title),
        Paragraph("MedAI радиология ассистенті", subtitle),
    ]

    metadata = [
        [
            _pdf_para("Зерттеу / Исследование", small_caps),
            _pdf_para(study.accession_number, normal),
            _pdf_para("Күні / Дата", small_caps),
            _pdf_para(_format_dt(report.confirmed_at), normal),
        ],
        [
            _pdf_para("Пациент / Пациент", small_caps),
            _pdf_para(study.patient_code, normal),
            _pdf_para("Түрі / Тип", small_caps),
            _pdf_para(study.study_type, normal),
        ],
        [
            _pdf_para("Дәрігер / Врач", small_caps),
            _pdf_para(doctor.full_name, normal),
            _pdf_para("AI мәртебесі / AI статус", small_caps),
            _pdf_para(chip_text, normal),
        ],
    ]
    table = Table(metadata, colWidths=[35 * mm, 57 * mm, 35 * mm, 57 * mm], hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7fafb")),
                ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#d7e1e6")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e4ebef")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TEXTCOLOR", (3, 2), (3, 2), colors.HexColor(chip_color)),
            ]
        )
    )
    story.extend([table, Spacer(1, 6)])

    if study.clinical_note:
        story.append(Paragraph("Клиникалық жазба / Клиническая заметка", section))
        kk_note = translate_text(study.clinical_note, "kk") or study.clinical_note
        note_box = Table([[_pdf_para(kk_note, normal)]], colWidths=[184 * mm])
        note_box.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fbfdfe")),
                    ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#dce6eb")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([note_box, Spacer(1, 6)])

    image_path = _latest_image_path(study)
    if image_path:
        story.append(Paragraph("Диагностикалық сурет / Диагностический снимок", section))
        image = _pdf_image(image_path, 92 * mm, 68 * mm)
        image_frame = Table([[image]], colWidths=[102 * mm], hAlign="CENTER")
        image_frame.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                    ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#d5e0e6")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        story.extend([image_frame, Spacer(1, 6)])

    conclusion_flow: list[Any] = []
    for heading, body in _clean_final_text(_report_text_for_lang(study, report, analysis, "kk"), study):
        conclusion_flow.append(Paragraph(heading, section))
        for paragraph in _plain_paragraphs(body):
            conclusion_flow.append(_pdf_para(paragraph, normal))
            conclusion_flow.append(Spacer(1, 3))
    story.append(KeepTogether(conclusion_flow))

    warning = Table([[_pdf_para(DISCLAIMER, muted)]], colWidths=[184 * mm])
    warning.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fff8ed")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#ead2a8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([warning, Spacer(1, 7)])

    photo = _pdf_image(_doctor_photo(doctor, export_dir), 18 * mm, 22 * mm)
    stamp_path = _stamp_path()
    stamp = _pdf_image(stamp_path, 22 * mm, 22 * mm) if stamp_path else _pdf_para("Мөр / Печать", muted)
    sign_text = _pdf_para(
        f"Дәрігер / Врач: {doctor.full_name}\n"
        f"Қолы / Подпись: ______________________________\n"
        f"Күні / Дата: {_format_dt(report.confirmed_at)}",
        normal,
    )
    sign_table = Table([[photo, sign_text, stamp]], colWidths=[26 * mm, 124 * mm, 34 * mm], hAlign="CENTER")
    sign_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fbfdfe")),
                ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#d5e0e6")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e4ebef")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, 0), "CENTER"),
                ("ALIGN", (2, 0), (2, 0), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(sign_table)

    ru_chip_text, ru_chip_color = _status_chip(analysis, "ru")
    story.append(PageBreak())
    story.extend([Paragraph("Медицинское заключение", title), Paragraph("Радиологический ассистент MedAI", subtitle)])
    ru_metadata = [
        [_pdf_para("Исследование", small_caps), _pdf_para(study.accession_number, normal), _pdf_para("Дата", small_caps), _pdf_para(_format_dt(report.confirmed_at), normal)],
        [_pdf_para("Пациент", small_caps), _pdf_para(study.patient_code, normal), _pdf_para("Тип", small_caps), _pdf_para(study.study_type, normal)],
        [_pdf_para("Врач", small_caps), _pdf_para(doctor.full_name, normal), _pdf_para("AI статус", small_caps), _pdf_para(ru_chip_text, normal)],
    ]
    ru_table = Table(ru_metadata, colWidths=[35 * mm, 57 * mm, 35 * mm, 57 * mm], hAlign="CENTER")
    ru_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7fafb")),
                ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#d7e1e6")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e4ebef")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TEXTCOLOR", (3, 2), (3, 2), colors.HexColor(ru_chip_color)),
            ]
        )
    )
    story.extend([ru_table, Spacer(1, 6)])

    if study.clinical_note:
        ru_note = translate_text(study.clinical_note, "ru") or study.clinical_note
        story.append(Paragraph("Клиническая заметка", section))
        ru_note_box = Table([[_pdf_para(ru_note, normal)]], colWidths=[184 * mm])
        ru_note_box.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fbfdfe")),
                    ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#dce6eb")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([ru_note_box, Spacer(1, 6)])

    image_path = _latest_image_path(study)
    if image_path:
        story.append(Paragraph("Диагностический снимок", section))
        ru_image = _pdf_image(image_path, 92 * mm, 68 * mm)
        ru_image_frame = Table([[ru_image]], colWidths=[102 * mm], hAlign="CENTER")
        ru_image_frame.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                    ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#d5e0e6")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        story.extend([ru_image_frame, Spacer(1, 6)])

    ru_conclusion_flow: list[Any] = []
    for heading, body in _clean_final_text(_report_text_for_lang(study, report, analysis, "ru"), study):
        ru_conclusion_flow.append(Paragraph(heading, section))
        for paragraph in _plain_paragraphs(body):
            ru_conclusion_flow.append(_pdf_para(paragraph, normal))
            ru_conclusion_flow.append(Spacer(1, 3))
    story.append(KeepTogether(ru_conclusion_flow))

    ru_warning = Table([[_pdf_para("Результат AI является предварительной подсказкой. Окончательное решение принимает врач.", muted)]], colWidths=[184 * mm])
    ru_warning.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fff8ed")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#ead2a8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([ru_warning, Spacer(1, 7)])

    ru_photo = _pdf_image(_doctor_photo(doctor, export_dir), 18 * mm, 22 * mm)
    ru_stamp_path = _stamp_path()
    ru_stamp = _pdf_image(ru_stamp_path, 22 * mm, 22 * mm) if ru_stamp_path else _pdf_para("Печать", muted)
    ru_sign_text = _pdf_para(
        f"Врач: {doctor.full_name}\n"
        f"Подпись: ______________________________\n"
        f"Дата: {_format_dt(report.confirmed_at)}",
        normal,
    )
    ru_sign_table = Table([[ru_photo, ru_sign_text, ru_stamp]], colWidths=[26 * mm, 124 * mm, 34 * mm], hAlign="CENTER")
    ru_sign_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fbfdfe")),
                ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#d5e0e6")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e4ebef")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, 0), "CENTER"),
                ("ALIGN", (2, 0), (2, 0), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(ru_sign_table)

    doc.build(story)
    return path


def _docx_set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    from docx.shared import Pt, RGBColor

    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _docx_paragraph(cell_or_doc, text: str = "", size: int = 10, bold: bool = False, color: str | None = None):
    paragraph = cell_or_doc.add_paragraph()
    run = paragraph.add_run(text)
    _docx_set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def _docx_shade(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _docx_cell_text(cell, text: str, size: int = 9, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _docx_set_run_font(run, size=size, bold=bold, color=color)


def export_report_docx(study: Study, report: Report, doctor: User, lang: str | None = "ru") -> Path:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    export_dir = ensure_export_dir()
    path = export_dir / f"study-{study.id}-report.docx"
    lang = "kk"
    analysis = _analysis_for_study(study)
    chip_text, _ = _status_chip(analysis, lang)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Медициналық қорытынды")
    _docx_set_run_font(run, size=14, bold=True, color="0B1F2A")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("MedAI радиология ассистенті")
    _docx_set_run_font(subrun, size=9, color="667985")

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    rows = [
        ("Зерттеу / Исследование", study.accession_number, "Күні / Дата", _format_dt(report.confirmed_at)),
        ("Пациент / Пациент", study.patient_code, "Түрі / Тип", study.study_type),
        ("Дәрігер / Врач", doctor.full_name, "AI мәртебесі / AI статус", chip_text),
    ]
    for row, values in zip(table.rows, rows):
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _docx_shade(cell, "F7FAFB")
            _docx_cell_text(cell, str(value), size=8 if index in (0, 2) else 9, bold=index in (1, 3), color="64748B" if index in (0, 2) else "17232B")

    if study.clinical_note:
        _docx_paragraph(doc, "Клиникалық жазба / Клиническая заметка", size=10, bold=True, color="0B6476")
        note_table = doc.add_table(rows=1, cols=1)
        note_table.style = "Table Grid"
        _docx_shade(note_table.cell(0, 0), "FBFDFE")
        _docx_cell_text(note_table.cell(0, 0), translate_text(study.clinical_note, "kk") or study.clinical_note, size=9, color="17232B")

    image_path = _latest_image_path(study)
    if image_path:
        _docx_paragraph(doc, "Диагностикалық сурет / Диагностический снимок", size=10, bold=True, color="0B6476")
        image_p = doc.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.add_run().add_picture(str(image_path), width=Inches(3.55))

    for heading, body in _clean_final_text(_report_text_for_lang(study, report, analysis, "kk"), study):
        _docx_paragraph(doc, heading, size=10, bold=True, color="0B6476")
        for paragraph_text in _plain_paragraphs(body):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(paragraph_text)
            _docx_set_run_font(run, size=9, color="17232B")

    warning = doc.add_table(rows=1, cols=1)
    warning.style = "Table Grid"
    _docx_shade(warning.cell(0, 0), "FFF8ED")
    _docx_cell_text(warning.cell(0, 0), DISCLAIMER, size=8, color="667985")

    sign = doc.add_table(rows=1, cols=3)
    sign.style = "Table Grid"
    cells = sign.rows[0].cells
    widths = [0.82, 4.55, 1.0]
    for cell, width in zip(cells, widths):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _docx_shade(cell, "FBFDFE")

    photo_paragraph = cells[0].paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_paragraph.add_run().add_picture(str(_doctor_photo(doctor, export_dir)), width=Inches(0.62))

    _docx_cell_text(
        cells[1],
        (
            f"Дәрігер / Врач: {doctor.full_name}\n"
            "Қолы / Подпись: ______________________________\n"
            f"Күні / Дата: {_format_dt(report.confirmed_at)}"
        ),
        size=8,
        color="17232B",
    )
    stamp_path = _stamp_path()
    stamp_paragraph = cells[2].paragraphs[0]
    stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if stamp_path:
        stamp_paragraph.add_run().add_picture(str(stamp_path), width=Inches(0.72))
    else:
        _docx_cell_text(cells[2], "Мөр / Печать", size=8, color="667985")

    doc.add_page_break()
    ru_chip_text, _ = _status_chip(analysis, "ru")

    ru_title = doc.add_paragraph()
    ru_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ru_title_run = ru_title.add_run("Медицинское заключение")
    _docx_set_run_font(ru_title_run, size=14, bold=True, color="0B1F2A")
    ru_subtitle = doc.add_paragraph()
    ru_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ru_subtitle_run = ru_subtitle.add_run("Радиологический ассистент MedAI")
    _docx_set_run_font(ru_subtitle_run, size=9, color="667985")

    ru_table = doc.add_table(rows=3, cols=4)
    ru_table.style = "Table Grid"
    ru_rows = [
        ("Исследование", study.accession_number, "Дата", _format_dt(report.confirmed_at)),
        ("Пациент", study.patient_code, "Тип", study.study_type),
        ("Врач", doctor.full_name, "AI статус", ru_chip_text),
    ]
    for row, values in zip(ru_table.rows, ru_rows):
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _docx_shade(cell, "F7FAFB")
            _docx_cell_text(cell, str(value), size=8 if index in (0, 2) else 9, bold=index in (1, 3), color="64748B" if index in (0, 2) else "17232B")

    if study.clinical_note:
        _docx_paragraph(doc, "Клиническая заметка", size=10, bold=True, color="0B6476")
        ru_note_table = doc.add_table(rows=1, cols=1)
        ru_note_table.style = "Table Grid"
        _docx_shade(ru_note_table.cell(0, 0), "FBFDFE")
        _docx_cell_text(ru_note_table.cell(0, 0), translate_text(study.clinical_note, "ru") or study.clinical_note, size=9, color="17232B")

    image_path = _latest_image_path(study)
    if image_path:
        _docx_paragraph(doc, "Диагностический снимок", size=10, bold=True, color="0B6476")
        ru_image_p = doc.add_paragraph()
        ru_image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ru_image_p.add_run().add_picture(str(image_path), width=Inches(3.55))

    for heading, body in _clean_final_text(_report_text_for_lang(study, report, analysis, "ru"), study):
        _docx_paragraph(doc, heading, size=10, bold=True, color="0B6476")
        for paragraph_text in _plain_paragraphs(body):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(paragraph_text)
            _docx_set_run_font(run, size=9, color="17232B")

    ru_warning = doc.add_table(rows=1, cols=1)
    ru_warning.style = "Table Grid"
    _docx_shade(ru_warning.cell(0, 0), "FFF8ED")
    _docx_cell_text(ru_warning.cell(0, 0), "Результат AI является предварительной подсказкой. Окончательное решение принимает врач.", size=8, color="667985")

    ru_sign = doc.add_table(rows=1, cols=3)
    ru_sign.style = "Table Grid"
    ru_cells = ru_sign.rows[0].cells
    for cell, width in zip(ru_cells, widths):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _docx_shade(cell, "FBFDFE")

    ru_photo_paragraph = ru_cells[0].paragraphs[0]
    ru_photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ru_photo_paragraph.add_run().add_picture(str(_doctor_photo(doctor, export_dir)), width=Inches(0.62))
    _docx_cell_text(
        ru_cells[1],
        f"Врач: {doctor.full_name}\nПодпись: ______________________________\nДата: {_format_dt(report.confirmed_at)}",
        size=8,
        color="17232B",
    )
    ru_stamp_paragraph = ru_cells[2].paragraphs[0]
    ru_stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if stamp_path:
        ru_stamp_paragraph.add_run().add_picture(str(stamp_path), width=Inches(0.72))
    else:
        _docx_cell_text(ru_cells[2], "Печать", size=8, color="667985")

    doc.save(path)
    return path
